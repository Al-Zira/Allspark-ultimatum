# File: frontend/src/main.py
import streamlit as st
import requests
import json
from typing import List, Dict
import os
import tempfile
import docx
import io
from fpdf import FPDF

# API endpoint
API_BASE_URL = "http://localhost:8080"

class MultiDocRedactionFrontend:
    def __init__(self):
        self._initialize_session_state()

    def _initialize_session_state(self):
        """Initialize session state variables"""
        defaults = {
            'custom_redactions': [],
            'contextual_matches': [],
            'document_text': "",
            'text_to_redact': "",
            'reason': "",
            'ai_suggestions': [],
            'has_generated_suggestions': False,
            'current_text': "",
            'current_type': "",
            'current_reason': "",
            'temp_file_path': None
        }
        for key, default_value in defaults.items():
            if key not in st.session_state:
                st.session_state[key] = default_value

    def _handle_file_upload(self, uploaded_file):
        """Handle file upload and text extraction"""
        if uploaded_file:
            try:
                # Create form data
                files = {
                    'file': uploaded_file
                }
                
                # Make the request
                response = requests.post(
                    f"{API_BASE_URL}/upload",
                    files=files,
                    timeout=30
                )
                
                try:
                    resp_json = response.json()
                    if response.status_code == 200:
                        st.session_state.document_text = resp_json['text']
                        st.session_state.temp_file_path = resp_json['file_path']
                        return True
                    else:
                        st.error(f"Error uploading file: {resp_json.get('detail', 'Unknown error')}")
                except Exception as e:
                    st.error(f"Error parsing response: {str(e)}")
                return False
            except Exception as e:
                st.error(f"Error during file upload: {str(e)}")
                return False

    def _handle_ai_suggestions(self):
        """Handle AI suggestion generation and display"""
        if st.button("Get AI Suggestions"):
            with st.spinner("Analyzing document for sensitive information..."):
                response = requests.post(
                    f"{API_BASE_URL}/ai-suggestions",  # Remove trailing slash
                    params={"file_path": st.session_state.temp_file_path}
                )
                if response.status_code == 200:
                    st.session_state.ai_suggestions = response.json()['suggestions']
                    st.session_state.has_generated_suggestions = True
                    st.rerun()
                else:
                    st.error(f"Error getting AI suggestions: {response.text}")

    def _display_ai_suggestions(self):
        """Display AI suggestions with accept buttons"""
        st.subheader("AI-Suggested Redactions")
        for i, suggestion in enumerate(st.session_state.ai_suggestions):
            with st.expander(f"Suggestion #{i+1}: {suggestion['type']}"):
                st.markdown(f"""
                - **Text to Redact:** `{suggestion['text']}`
                - **Type:** {suggestion['type']}
                - **Confidence:** {suggestion['confidence']}%
                - **Reason:** {suggestion['reason']}
                """)
                if st.button("Accept Suggestion", key=f"accept_{i}"):
                    self._handle_suggestion_acceptance(suggestion)

    def _handle_suggestion_acceptance(self, suggestion: dict):
        """Handle the acceptance of an AI suggestion"""
        if suggestion['text'] not in [r['text'] for r in st.session_state.custom_redactions]:
            st.session_state.custom_redactions.append(suggestion)
            st.success("Suggestion added to redaction list!")
        else:
            st.warning("This text is already in the redaction list!")

    def _handle_custom_redaction(self):
        """Handle custom redaction input and processing"""
        text_to_redact = st.text_input("Enter text to redact:")
        redaction_type = st.selectbox(
            "Redaction Type:",
            ["PII", "CREDENTIALS", "FINANCIAL", "CUSTOM"]
        )
        reason = st.text_input("Reason for redaction (optional):")
        
        if st.button("Add Redaction") and text_to_redact:
            try:
                # Make the request
                response = requests.post(
                    f"{API_BASE_URL}/analyze-context",
                    params={
                        "document_text": st.session_state.document_text,
                        "text": text_to_redact,
                        "type": redaction_type
                    }
                )
                
                if response.status_code == 200:
                    matches = response.json()['matches']
                    if matches:
                        st.session_state.contextual_matches = matches
                        st.session_state.current_text = text_to_redact
                        st.session_state.current_type = redaction_type
                        st.session_state.current_reason = reason
                        st.success(f"Found {len(matches)} matches for '{text_to_redact}'")
                    else:
                        st.warning(f"No matches found for '{text_to_redact}'")
                else:
                    st.error(f"Error analyzing context: {response.json().get('detail', 'Unknown error')}")
            except Exception as e:
                st.error(f"Error during context analysis: {str(e)}")

    def _display_contextual_matches(self):
        """Display contextual matches with accept buttons"""
        if st.session_state.contextual_matches:
            st.subheader("Contextual Matches Found:")
            for idx, match in enumerate(st.session_state.contextual_matches):
                with st.expander(f"Match #{idx + 1}: {match['text']}"):
                    st.markdown(f"""
                    - **Found Text:** `{match['text']}`
                    - **Context:** `{match['context']}`
                    - **Confidence:** {match['confidence']}%
                    - **Reason:** {match['reason']}
                    """)
                    
                    if st.button("Add this match", key=f"add_match_{idx}"):
                        new_redaction = {
                            'text': match['text'],
                            'type': st.session_state.current_type,
                            'confidence': match['confidence'],
                            'reason': st.session_state.current_reason or match['reason']
                        }
                        
                        if new_redaction not in st.session_state.custom_redactions:
                            st.session_state.custom_redactions.append(new_redaction)
                            st.success(f"Added redaction for: {match['text']}")
                            st.rerun()

    def _display_current_redactions(self):
        """Display current redactions with remove buttons"""
        if st.session_state.custom_redactions:
            st.subheader("Current Redactions")
            for idx, redaction in enumerate(st.session_state.custom_redactions):
                with st.expander(f"Redaction #{idx + 1}: {redaction['text']}"):
                    st.markdown(f"""
                    - **Text:** `{redaction['text']}`
                    - **Type:** {redaction['type']}
                    - **Confidence:** {redaction['confidence']}%
                    - **Reason:** {redaction['reason']}
                    """)
                    if st.button("Remove", key=f"remove_{idx}"):
                        st.session_state.custom_redactions.pop(idx)
                        st.rerun()

    def _apply_redactions(self):
        """Apply redactions and handle response"""
        if st.button("Apply Redactions") and st.session_state.custom_redactions:
            with st.spinner("Applying redactions..."):
                response = requests.post(
                    f"{API_BASE_URL}/apply-redactions",  # Remove trailing slash
                    json={
                        "redactions": st.session_state.custom_redactions,
                        "file_path": st.session_state.temp_file_path
                    }
                )
                if response.status_code == 200:
                    data = response.json()
                    st.success("Redactions applied successfully!")
                    
                    # Download buttons for redacted file and report
                    with open(data['redacted_file_path'], 'rb') as f:
                        st.download_button(
                            "Download Redacted Document",
                            f,
                            file_name=f"redacted_{os.path.basename(st.session_state.temp_file_path)}"
                        )
                    
                    st.download_button(
                        "Download Redaction Report",
                        json.dumps(data['report'], indent=2),
                        file_name="redaction_report.json",
                        mime="application/json"
                    )
                else:
                    st.error(f"Error applying redactions: {response.text}")

    def show_sample_files(self):
        """Display and handle sample file generation"""
        st.subheader("Sample Test Files")
        response = requests.get(f"{API_BASE_URL}/sample-text")  # Remove trailing slash
        if response.status_code == 200:
            sample_text = response.json()['text']
            
            # Create download buttons for each format
            col1, col2, col3 = st.columns(3)
            
            with col1:
                st.download_button(
                    label="📄 Download TXT Sample",
                    data=sample_text,
                    file_name="sample_document.txt",
                    mime="text/plain"
                )
            
            with col2:
                # Create a simple DOCX in memory
                doc = docx.Document()
                doc.add_heading('Sample Document', 0)
                for paragraph in sample_text.split('\n'):
                    doc.add_paragraph(paragraph)
                
                # Save to bytes
                docx_bytes = io.BytesIO()
                doc.save(docx_bytes)
                docx_bytes.seek(0)
                
                st.download_button(
                    label="📘 Download DOCX Sample",
                    data=docx_bytes,
                    file_name="sample_document.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            
            with col3:
                # Create PDF in memory
                pdf = FPDF()
                pdf.add_page()
                pdf.set_font("Arial", size=12)
                for line in sample_text.split('\n'):
                    pdf.cell(0, 10, txt=line, ln=True)
                
                pdf_bytes = pdf.output(dest='S').encode('latin-1')
                
                st.download_button(
                    label="📕 Download PDF Sample",
                    data=pdf_bytes,
                    file_name="sample_document.pdf",
                    mime="application/pdf"
                )

    def _redact_text(self, file_path: str, suggestions: List[Dict], output_path: str) -> str:
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Sort suggestions by length and position to handle overlapping redactions
            suggestions.sort(key=lambda x: (-len(x['text']), x['text'].lower()))
            
            # Create a list of all positions to redact
            redactions = []
            for suggestion in suggestions:
                text = suggestion['text']
                start = 0
                while True:
                    pos = content.find(text, start)
                    if pos == -1:
                        break
                    redactions.append((pos, pos + len(text), '' * len(text)))
                    start = pos + 1
            
            # Sort redactions by start position in reverse order
            redactions.sort(key=lambda x: x[0], reverse=True)
            
            # Apply redactions from end to start to◼️ maintain correct positions
            content_chars = list(content)
            for start, end, replacement in redactions:
                content_chars[start:end] = replacement
            
            # Convert back to string
            redacted_content = ''.join(content_chars)
            
            # Write redacted content
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(redacted_content)
            
            return output_path
            
        except Exception as e:
            raise RuntimeError(f"Text redaction failed: {e}")

    def run(self):
        """Main application flow"""
        st.title("MultiDoc Redaction Assistant")
        
        self.show_sample_files()
        
        uploaded_file = st.file_uploader(
            "Upload a document",
            type=["pdf", "docx", "txt"],
            help="Supported formats: PDF, DOCX, TXT"
        )
        
        if uploaded_file:
            if self._handle_file_upload(uploaded_file):
                st.subheader("Document Preview")
                st.text(st.session_state.document_text)
                
                self._handle_ai_suggestions()
                
                if st.session_state.has_generated_suggestions and st.session_state.ai_suggestions:
                    self._display_ai_suggestions()
                elif st.session_state.has_generated_suggestions:
                    st.info("No sensitive information detected by AI.")
                
                st.subheader("Custom Redaction")
                self._handle_custom_redaction()
                
                # Display contextual matches
                self._display_contextual_matches()
                
                # Display current redactions
                self._display_current_redactions()
                
                # Apply redactions button
                if st.session_state.custom_redactions:
                    self._apply_redactions()

if __name__ == "__main__":
    app = MultiDocRedactionFrontend()
    app.run()