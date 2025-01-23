from fastapi import FastAPI, UploadFile, File, HTTPException, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.middleware.trustedhost import TrustedHostMiddleware
from pydantic import BaseModel
from typing import List, Dict, Optional
import tempfile
import os
import logging
from PyPDF2 import PdfReader, PdfWriter
import docx
import re
from reportlab.pdfgen import canvas
from reportlab.lib.colors import black
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from io import BytesIO
import fitz  # PyMuPDF
import google.generativeai as genai
from dotenv import load_dotenv
from functools import lru_cache
import asyncio
from concurrent.futures import ThreadPoolExecutor
import hashlib
import time
import aiofiles

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Constants
MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB
CHUNK_SIZE = 1024 * 1024  # 1MB
ALLOWED_EXTENSIONS = {'.pdf', '.docx', '.txt'}
MAX_REQUESTS_PER_MINUTE = 120  # Increased from 60

# Pre-compile regex patterns
PATTERNS = {
    'email': (re.compile(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'), 'EMAIL'),
    'phone': (re.compile(r'\+?\d{1,3}[-.\s]?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}'), 'PHONE'),
    'ssn': (re.compile(r'\b\d{3}[-]?\d{2}[-]?\d{4}\b'), 'SSN'),
    'credit_card': (re.compile(r'\b\d{4}[\s-]?\d{4}[\s-]?\d{4}[\s-]?\d{4}\b'), 'CREDIT_CARD'),
    'address': (re.compile(r'\b\d+\s+[A-Za-z]+\s+(?:Street|St|Avenue|Ave|Road|Rd|Boulevard|Blvd|Lane|Ln|Drive|Dr|Court|Ct|Circle|Cir|Way)[,.]?\s+[A-Za-z]+(?:[,.]?\s*[A-Za-z]+)?\b'), 'ADDRESS'),
    'name': (re.compile(r'\b[A-Z][a-z]+\s+[A-Z][a-z]+\b'), 'NAME')
}

# Rate limiting
request_counts: Dict[str, List[float]] = {}

app = FastAPI(debug=False)  # Disable debug mode in production

# Add security middleware
app.add_middleware(
    TrustedHostMiddleware,
    allowed_hosts=["localhost", "127.0.0.1"]
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:8501"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Load environment variables
load_dotenv()

# Configure Gemini AI
genai.configure(api_key=os.getenv('GOOGLE_API_KEY'))
model = genai.GenerativeModel('gemini-pro')

# Initialize thread pool
thread_pool = ThreadPoolExecutor(max_workers=4)

class RedactionItem(BaseModel):
    text: str
    type: str
    confidence: float
    reason: str

class RedactionRequest(BaseModel):
    redactions: List[RedactionItem]
    file_path: str

async def rate_limit(request: Request):
    client_ip = request.client.host
    now = time.time()
    
    if client_ip not in request_counts:
        request_counts[client_ip] = []
    
    # Remove old requests
    request_counts[client_ip] = [t for t in request_counts[client_ip] if now - t < 60]
    
    if len(request_counts[client_ip]) >= MAX_REQUESTS_PER_MINUTE:
        raise HTTPException(status_code=429, detail="Too many requests")
    
    request_counts[client_ip].append(now)

@lru_cache(maxsize=100)
def find_sensitive_info(text: str) -> List[dict]:
    """Find sensitive information in text with caching"""
    findings = []
    for pattern_name, (pattern, type_) in PATTERNS.items():
        matches = pattern.finditer(text)
        for match in matches:
            findings.append({
                "text": match.group(),
                "type": "PII",
                "confidence": 0.95,
                "reason": f"Contains {type_}"
            })
    
    return findings

async def process_file_chunk(chunk: bytes) -> str:
    """Process a chunk of file data"""
    return chunk.decode('utf-8', errors='ignore')

async def extract_text_from_file(file_path: str) -> str:
    """Extract text from file using async processing for large files"""
    file_ext = os.path.splitext(file_path)[1].lower()
    
    if file_ext not in ALLOWED_EXTENSIONS:
        raise ValueError(f"Unsupported file type: {file_ext}")
    
    try:
        if file_ext == '.pdf':
            loop = asyncio.get_event_loop()
            with open(file_path, 'rb') as file:
                pdf = await loop.run_in_executor(thread_pool, PdfReader, file)
                text = ""
                for page in pdf.pages:
                    page_text = await loop.run_in_executor(thread_pool, page.extract_text)
                    text += page_text + "\n"
                return text
        elif file_ext == '.docx':
            loop = asyncio.get_event_loop()
            doc = await loop.run_in_executor(thread_pool, docx.Document, file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            return text
        elif file_ext == '.txt':
            text = ""
            async with aiofiles.open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                chunk = await file.read(CHUNK_SIZE)
                while chunk:
                    text += await process_file_chunk(chunk.encode())
                    chunk = await file.read(CHUNK_SIZE)
            return text
    except Exception as e:
        logger.error(f"Error extracting text: {str(e)}")
        raise

def cleanup_temp_files():
    """Clean up temporary files older than 1 hour"""
    temp_dir = tempfile.gettempdir()
    current_time = time.time()
    
    for filename in os.listdir(temp_dir):
        if filename.startswith(('upload_', 'redacted_')):
            filepath = os.path.join(temp_dir, filename)
            if current_time - os.path.getctime(filepath) > 3600:  # 1 hour
                try:
                    os.remove(filepath)
                    logger.info(f"Cleaned up temp file: {filepath}")
                except Exception as e:
                    logger.error(f"Error cleaning up {filepath}: {str(e)}")

@app.middleware("http")
async def add_security_headers(request: Request, call_next):
    """Add security headers to all responses"""
    response = await call_next(request)
    response.headers["X-Content-Type-Options"] = "nosniff"
    response.headers["X-Frame-Options"] = "DENY"
    response.headers["X-XSS-Protection"] = "1; mode=block"
    return response

@app.get("/")
async def root(request: Request):
    await rate_limit(request)
    logger.info("Root endpoint called")
    return {"message": "API is running"}

@app.get("/test")
async def test(request: Request):
    await rate_limit(request)
    logger.info("Test endpoint called")
    return {"message": "API is working"}

@app.get("/sample-text")
async def get_sample_text():
    """Get sample text for testing"""
    return {"text": """John Doe lives at 123 Elm Street, Springfield.
His email address is john.doe@example.com, and his phone number is +1-555-123-4567.
Credit Card Number: 4111 1111 1111 1111. Social Security Number: 123-45-6789.

For work-related matters, contact him at work.email@example.com or at (555) 987-6543.
He frequently shops online and uses PayPal with the email paypal.john.doe@example.com."""}

@app.post("/upload")
async def upload_file(request: Request, file: UploadFile = File(...)):
    await rate_limit(request)
    try:
        logger.info(f"Upload endpoint called with file: {file.filename}")
        
        # Validate file extension
        file_ext = os.path.splitext(file.filename)[1].lower()
        if file_ext not in ALLOWED_EXTENSIONS:
            raise HTTPException(status_code=400, detail=f"Unsupported file type: {file_ext}")
        
        # Read file in chunks
        content = BytesIO()
        size = 0
        while chunk := await file.read(CHUNK_SIZE):
            size += len(chunk)
            if size > MAX_FILE_SIZE:
                raise HTTPException(status_code=400, detail="File too large")
            content.write(chunk)
        
        logger.info(f"File content read, size: {size} bytes")
        
        # Create temp directory if it doesn't exist
        temp_dir = tempfile.gettempdir()
        os.makedirs(temp_dir, exist_ok=True)
        
        # Generate unique filename using content hash
        content_hash = hashlib.md5(content.getvalue()).hexdigest()
        temp_file_path = os.path.join(temp_dir, f"upload_{content_hash}_{file.filename}")
        
        # Save file
        with open(temp_file_path, "wb") as f:
            f.write(content.getvalue())
        logger.info(f"File saved to: {temp_file_path}")
        
        # Extract text from the file
        extracted_text = await extract_text_from_file(temp_file_path)
        logger.info(f"Text extracted, length: {len(extracted_text)} characters")
        
        # Schedule cleanup
        asyncio.create_task(asyncio.to_thread(cleanup_temp_files))
        
        return {
            "message": f"Successfully received file: {file.filename}",
            "size": size,
            "file_path": temp_file_path,
            "text": extracted_text
        }
    except Exception as e:
        logger.error(f"Error processing upload: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/ai-suggestions")
async def get_ai_suggestions(request: Request, file_path: str):
    await rate_limit(request)
    try:
        logger.info(f"AI suggestions requested for file: {file_path}")
        
        # Validate file path
        if not os.path.exists(file_path):
            raise HTTPException(status_code=404, detail="File not found")
        
        # Read the file content
        text = await extract_text_from_file(file_path)
        logger.info(f"Text extracted for AI analysis, length: {len(text)} characters")
        
        # Find sensitive information using cached function
        suggestions = find_sensitive_info(text)
        logger.info(f"Found {len(suggestions)} sensitive items")
        
        return {"suggestions": suggestions}
    except Exception as e:
        logger.error(f"Error getting AI suggestions: {str(e)}")
        raise HTTPException(status_code=400, detail=str(e))

@app.post("/apply-redactions")
async def apply_redactions(request: Request, redaction_request: RedactionRequest):
    await rate_limit(request)
    try:
        logger.info(f"Applying {len(redaction_request.redactions)} redactions to file: {redaction_request.file_path}")
        
        # Validate file path
        if not os.path.exists(redaction_request.file_path):
            raise HTTPException(status_code=404, detail="File not found")
        
        # Generate unique output path using content hash
        content_hash = hashlib.md5(open(redaction_request.file_path, 'rb').read()).hexdigest()
        output_path = os.path.join(
            tempfile.gettempdir(),
            f"redacted_{content_hash}_{os.path.basename(redaction_request.file_path)}"
        )
        
        # Process file based on extension
        file_ext = os.path.splitext(redaction_request.file_path)[1].lower()
        if file_ext == '.pdf':
            # Use thread pool for PDF processing
            loop = asyncio.get_event_loop()
            await loop.run_in_executor(
                thread_pool,
                process_pdf_redactions,
                redaction_request.file_path,
                output_path,
                redaction_request.redactions
            )
        elif file_ext == '.docx':
            await process_docx_redactions(
                redaction_request.file_path,
                output_path,
                redaction_request.redactions
            )
        else:
            await process_text_redactions(
                redaction_request.file_path,
                output_path,
                redaction_request.redactions
            )
        
        # Schedule cleanup
        asyncio.create_task(asyncio.to_thread(cleanup_temp_files))
        
        # Generate report
        report = {
            "original_file": redaction_request.file_path,
            "redacted_file": output_path,
            "num_redactions": len(redaction_request.redactions),
            "redactions": [redaction.dict() for redaction in redaction_request.redactions]
        }
        
        return {
            "redacted_file_path": output_path,
            "report": report
        }
    except Exception as e:
        logger.error(f"Error applying redactions: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

def process_pdf_redactions(input_path: str, output_path: str, redactions: List[RedactionItem]):
    """Process PDF redactions in a separate thread"""
    doc = fitz.open(input_path)
    for page in doc:
        for redaction in redactions:
            instances = page.search_for(redaction.text)
            for inst in instances:
                page.add_redact_annot(inst)
        page.apply_redactions()
    doc.save(output_path)
    doc.close()

async def process_docx_redactions(input_path: str, output_path: str, redactions: List[RedactionItem]):
    """Process DOCX redactions"""
    loop = asyncio.get_event_loop()
    doc = await loop.run_in_executor(thread_pool, docx.Document, input_path)
    
    for paragraph in doc.paragraphs:
        current_text = paragraph.text
        paragraph.clear()
        
        for redaction in redactions:
            parts = current_text.split(redaction.text)
            for i, part in enumerate(parts):
                if part:
                    run = paragraph.add_run(part)
                if i < len(parts) - 1:
                    run = paragraph.add_run('█' * len(redaction.text))
                    run.font.color.rgb = docx.shared.RGBColor(0, 0, 0)
    
    await loop.run_in_executor(thread_pool, doc.save, output_path)

async def process_text_redactions(input_path: str, output_path: str, redactions: List[RedactionItem]):
    """Process text file redactions"""
    async with aiofiles.open(input_path, 'r', encoding='utf-8') as f:
        content = await f.read()
    
    redacted_text = content
    for redaction in redactions:
        redacted_text = redacted_text.replace(redaction.text, '█' * len(redaction.text))
    
    async with aiofiles.open(output_path, 'w', encoding='utf-8') as f:
        await f.write(redacted_text)

@app.post("/analyze-context")
async def analyze_context(request: Request, document_text: str = None, text: str = None, type: str = None):
    """Analyze contextual meaning using Gemini AI"""
    try:
        if not all([document_text, text, type]):
            raise HTTPException(status_code=400, detail="Missing required parameters")
            
        logger.info(f"Analyzing context for text: {text}")
        
        # Prepare the prompt for Gemini
        prompt = f"""
        You are an expert in identifying sensitive information in text. Analyze the following text and find all instances of information related to "{text}".
        Consider semantic meaning and context, not just exact matches.

        For example:
        - If searching for "email", find all email addresses
        - If searching for "address", find all physical addresses
        - If searching for "phone", find all phone numbers
        - If searching for "name", find all person names
        - If searching for "card", find all credit card numbers
        - If searching for "ssn" or "social", find all social security numbers

        Text to analyze:
        {document_text}

        For each match found, provide a JSON object with these fields:
        {{
            "text": "the exact matched text",
            "type": "{type}",
            "confidence": <confidence score between 0.75 and 0.99>,
            "reason": "detailed explanation of why this is a match"
        }}

        Format your response as a JSON array of these objects. Include ONLY the JSON array, no other text.
        Ensure each confidence score reflects the certainty of the match (higher for exact matches, lower for contextual matches).
        Provide detailed, specific reasons for each match.
        """

        # Get response from Gemini
        response = model.generate_content(prompt)
        
        # Parse the response
        try:
            # Extract the JSON array from the response
            response_text = response.text
            # Clean up the response to ensure it's valid JSON
            response_text = response_text.replace('```json', '').replace('```', '').strip()
            matches = eval(response_text)  # Using eval since the response is a Python list/dict format
            
            # Add context and format confidence scores
            for match in matches:
                # Find the match in the original text
                start_idx = document_text.find(match['text'])
                if start_idx != -1:
                    # Get context around the match
                    start = max(0, start_idx - 30)
                    end = min(len(document_text), start_idx + len(match['text']) + 30)
                    match['context'] = document_text[start:end]
                
                # Format confidence score as percentage
                match['confidence'] = round(float(match['confidence']) * 100, 1)
        
        except Exception as e:
            logger.error(f"Error parsing Gemini response: {str(e)}")
            # Fallback to basic pattern matching if Gemini response can't be parsed
            matches = []
            # Define patterns for common sensitive information
            patterns = {
                'email': r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b',
                'phone': r'\+?\d{1,3}[-.\s]?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}',
                'ssn': r'\b\d{3}[-]?\d{2}[-]?\d{4}\b',
                'credit_card': r'\b\d{4}[\s-]?\d{4}[\s-]?\d{4}[\s-]?\d{4}\b',
                'address': r'\b\d+\s+[A-Za-z]+\s+(?:Street|St|Avenue|Ave|Road|Rd|Boulevard|Blvd|Lane|Ln|Drive|Dr|Court|Ct|Circle|Cir|Way)[,.]?\s+[A-Za-z]+(?:[,.]?\s*[A-Za-z]+)?\b',
                'name': r'\b[A-Z][a-z]+\s+[A-Z][a-z]+\b'
            }
            
            # Try to match based on semantic meaning
            pattern = None
            for key, pat in patterns.items():
                if text.lower() in key:
                    pattern = pat
                    break
            
            if not pattern:
                pattern = re.escape(text)
            
            # Find matches
            for match in re.finditer(pattern, document_text, re.IGNORECASE):
                start = max(0, match.start() - 30)
                end = min(len(document_text), match.end() + 30)
                context = document_text[start:end]
                
                matches.append({
                    "text": match.group(),
                    "type": type,
                    "confidence": 85.0,  # Lower confidence for fallback matches
                    "reason": f"Pattern match for {text}",
                    "context": context
                })
        
        logger.info(f"Found {len(matches)} matches for text: {text}")
        return {"matches": matches}
    except Exception as e:
        logger.error(f"Error analyzing context: {str(e)}")
        raise HTTPException(status_code=400, detail=str(e)) 